"""
PredictIQ Document Parser Service
===================================
Extracts raw text from PDF, DOCX, and TXT files stored in the
Neon PostgreSQL database. Returns clean text for NLP extraction.

Handles corrupt files gracefully — never crashes the API.
"""

import io
import re
import structlog
from typing import Optional

logger = structlog.get_logger()


class DocumentParser:
    """Handles text extraction from various document formats."""

    @staticmethod
    def parse(file_content: bytes, mime_type: str) -> dict:
        """
        Parse a document and extract its text content.

        Args:
            file_content: Raw bytes of the uploaded file.
            mime_type: MIME type of the file.

        Returns:
            Dictionary with raw_text, text_preview, word_count, page_count.

        Raises:
            ValueError: If file type is unsupported or content is unreadable.
        """
        try:
            if mime_type == "application/pdf":
                return DocumentParser._parse_pdf(file_content)
            elif mime_type in (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
                "application/msword",
            ):
                return DocumentParser._parse_docx(file_content)
            elif mime_type == "text/plain":
                return DocumentParser._parse_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
        except ValueError:
            raise
        except Exception as e:
            logger.error("document_parse_error", error=str(e), mime_type=mime_type)
            raise ValueError(f"Failed to parse document: {e}")

    @staticmethod
    def _parse_pdf(content: bytes) -> dict:
        """Extract text from PDF using pdfplumber."""
        import pdfplumber

        text_pages: list[str] = []
        page_count = 0

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_pages.append(page_text.strip())

                    # Also extract table content
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = " | ".join(
                                    str(cell).strip()
                                    for cell in row
                                    if cell and str(cell).strip()
                                )
                                if row_text:
                                    text_pages.append(row_text)

                    if i >= 49:  # Cap at 50 pages
                        text_pages.append(
                            "[...document truncated at 50 pages]"
                        )
                        break
        except Exception as e:
            logger.error("pdf_parse_error", error=str(e))
            raise ValueError(
                f"Could not parse PDF: {e}. "
                "Ensure it's a text-based PDF (not a scanned image)."
            )

        raw_text = DocumentParser._clean_text("\n\n".join(text_pages))

        if len(raw_text) < 50:
            raise ValueError(
                "PDF appears empty or is a scanned image. "
                "Please upload a text-based PDF."
            )

        return {
            "raw_text": raw_text,
            "text_preview": raw_text[:500],
            "word_count": len(raw_text.split()),
            "page_count": page_count,
        }

    @staticmethod
    def _parse_docx(content: bytes) -> dict:
        """Extract text from DOCX including paragraphs and tables."""
        from docx import Document

        text_parts: list[str] = []

        try:
            doc = Document(io.BytesIO(content))

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

        except Exception as e:
            logger.error("docx_parse_error", error=str(e))
            raise ValueError(f"Could not parse DOCX: {e}")

        raw_text = DocumentParser._clean_text("\n\n".join(text_parts))

        if len(raw_text) < 50:
            raise ValueError("DOCX appears empty or could not be read.")

        return {
            "raw_text": raw_text,
            "text_preview": raw_text[:500],
            "word_count": len(raw_text.split()),
            "page_count": None,
        }

    @staticmethod
    def _parse_txt(content: bytes) -> dict:
        """Decode plain text with multi-encoding fallback."""
        raw_text = ""
        for enc in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                raw_text = content.decode(enc)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if not raw_text:
            raise ValueError(
                "Could not decode text file "
                "(tried utf-8, utf-16, latin-1, cp1252)"
            )

        raw_text = DocumentParser._clean_text(raw_text)
        return {
            "raw_text": raw_text,
            "text_preview": raw_text[:500],
            "word_count": len(raw_text.split()),
            "page_count": None,
        }

    @staticmethod
    def _clean_text(text: str) -> str:
        """Normalize whitespace and remove excess blank lines."""
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        text = re.sub(r"\r\n", "\n", text)
        return text.strip()


document_parser = DocumentParser()
